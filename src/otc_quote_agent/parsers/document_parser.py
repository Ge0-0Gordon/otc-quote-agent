"""Simple text extraction for supported quote document formats."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from otc_quote_agent.schemas import SourceType


class DocumentParseError(ValueError):
    """Raised when a supported document cannot be parsed."""


class UnsupportedDocumentError(DocumentParseError):
    """Raised when the input extension is not supported."""


class ScannedPdfError(DocumentParseError):
    """Raised when a PDF contains no extractable text."""


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    source_file: str | None
    source_type: SourceType


class DocumentParser:
    """Extract plain text without reconstructing layout or running OCR."""

    SUPPORTED_EXTENSIONS = {
        ".txt": SourceType.TXT,
        ".md": SourceType.MD,
        ".docx": SourceType.DOCX,
        ".xlsx": SourceType.XLSX,
        ".pdf": SourceType.PDF,
    }

    def parse_text(self, text: str) -> ParsedDocument:
        cleaned = text.strip()
        if not cleaned:
            raise DocumentParseError("Pasted text is empty.")
        return ParsedDocument(
            text=cleaned,
            source_file=None,
            source_type=SourceType.PASTED_TEXT,
        )

    def parse_path(self, path: str | Path) -> ParsedDocument:
        file_path = Path(path)
        if not file_path.is_file():
            raise DocumentParseError(f"Input file does not exist: {file_path}")
        return self.parse_bytes(file_path.name, file_path.read_bytes())

    def parse_bytes(self, filename: str, content: bytes) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        source_type = self.SUPPORTED_EXTENSIONS.get(suffix)
        if source_type is None:
            supported = ", ".join(sorted(self.SUPPORTED_EXTENSIONS))
            raise UnsupportedDocumentError(
                f"Unsupported file type '{suffix or '(none)'}'. "
                f"Supported types: {supported}"
            )
        if not content:
            raise DocumentParseError(f"Input file is empty: {filename}")

        try:
            text = self._extract(source_type, content)
        except ScannedPdfError:
            raise
        except Exception as exc:
            raise DocumentParseError(f"Failed to parse {filename}: {exc}") from exc

        cleaned = "\n".join(line.rstrip() for line in text.splitlines()).strip()
        if not cleaned:
            if source_type is SourceType.PDF:
                raise ScannedPdfError(
                    "PDF contains no extractable text and may be scanned. "
                    "OCR is not supported in this version."
                )
            raise DocumentParseError(f"No extractable text found in {filename}.")

        return ParsedDocument(
            text=cleaned,
            source_file=filename,
            source_type=source_type,
        )

    def _extract(self, source_type: SourceType, content: bytes) -> str:
        if source_type in {SourceType.TXT, SourceType.MD}:
            return self._decode_text(content)
        if source_type is SourceType.DOCX:
            return self._extract_docx(content)
        if source_type is SourceType.XLSX:
            return self._extract_xlsx(content)
        if source_type is SourceType.PDF:
            return self._extract_pdf(content)
        raise UnsupportedDocumentError(f"Unsupported source type: {source_type}")

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8-sig", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("Text file must use UTF-8 or GB18030 encoding.")

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        document = Document(BytesIO(content))
        lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    lines.append(" | ".join(values))
        return "\n".join(lines)

    @staticmethod
    def _extract_xlsx(content: bytes) -> str:
        workbook = load_workbook(
            BytesIO(content),
            read_only=True,
            data_only=True,
        )
        lines: list[str] = []
        try:
            for sheet in workbook.worksheets:
                lines.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value).strip() for value in row if value is not None]
                    if values:
                        lines.append(" | ".join(values))
        finally:
            workbook.close()
        return "\n".join(lines)

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise DocumentParseError("Encrypted PDF files are not supported.")
        return "\n".join((page.extract_text() or "") for page in reader.pages)
