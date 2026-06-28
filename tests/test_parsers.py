from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from otc_quote_agent.parsers import (
    DocumentParser,
    InputLimitError,
    ScannedPdfError,
    UnsupportedDocumentError,
)
from otc_quote_agent.schemas import SourceType


@pytest.fixture
def parser() -> DocumentParser:
    return DocumentParser()


def test_parse_pasted_text(parser: DocumentParser) -> None:
    result = parser.parse_text("  Snowball inquiry  ")

    assert result.text == "Snowball inquiry"
    assert result.source_type is SourceType.PASTED_TEXT


@pytest.mark.parametrize("filename", ["quote.txt", "quote.md"])
def test_parse_plain_text(parser: DocumentParser, filename: str) -> None:
    result = parser.parse_bytes(filename, "中证1000雪球".encode())

    assert result.text == "中证1000雪球"


def test_parse_docx_paragraphs_and_tables(parser: DocumentParser) -> None:
    document = Document()
    document.add_paragraph("FCN quote")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Coupon"
    table.cell(0, 1).text = "12%"
    content = BytesIO()
    document.save(content)

    result = parser.parse_bytes("quote.docx", content.getvalue())

    assert "FCN quote" in result.text
    assert "Coupon | 12%" in result.text


def test_parse_xlsx_visible_cells(parser: DocumentParser) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Quote"
    sheet.append(["Product", "European Option"])
    content = BytesIO()
    workbook.save(content)

    result = parser.parse_bytes("quote.xlsx", content.getvalue())

    assert "[Sheet: Quote]" in result.text
    assert "Product | European Option" in result.text


def test_empty_text_pdf_is_reported_as_scanned(parser: DocumentParser) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    content = BytesIO()
    writer.write(content)

    with pytest.raises(ScannedPdfError, match="OCR is not supported"):
        parser.parse_bytes("scan.pdf", content.getvalue())


def test_unsupported_extension_is_rejected(parser: DocumentParser) -> None:
    with pytest.raises(UnsupportedDocumentError):
        parser.parse_bytes("quote.csv", b"product,snowball")


def test_parse_plain_text_email(parser: DocumentParser) -> None:
    content = (
        b"Subject: EURUSD option request\r\n"
        b"From: sales@example.com\r\n"
        b"To: trader@example.com\r\n"
        b"Content-Type: text/plain; charset=utf-8\r\n"
        b"\r\n"
        b"Please quote a 3-month EURUSD European call."
    )

    result = parser.parse_bytes("quote.eml", content)

    assert result.source_type is SourceType.EML
    assert "Subject: EURUSD option request" in result.text
    assert "Please quote a 3-month EURUSD European call." in result.text


def test_parse_html_email_when_plain_text_is_absent(parser: DocumentParser) -> None:
    content = (
        b"Subject: Snowball\r\n"
        b"Content-Type: text/html; charset=utf-8\r\n"
        b"\r\n"
        b"<p>Please quote <strong>Snowball</strong> at 15%.</p>"
    )

    result = parser.parse_bytes("quote.eml", content)

    assert "Please quote" in result.text
    assert "Snowball" in result.text
    assert "<strong>" not in result.text


def test_input_file_size_limit_is_explicit(parser: DocumentParser) -> None:
    parser.MAX_FILE_BYTES = 4

    with pytest.raises(InputLimitError, match="byte limit"):
        parser.parse_bytes("quote.txt", b"12345")


def test_extracted_text_length_limit_is_explicit(parser: DocumentParser) -> None:
    parser.MAX_TEXT_CHARS = 4

    with pytest.raises(InputLimitError, match="character limit"):
        parser.parse_text("12345")


def test_pdf_page_limit_is_explicit(parser: DocumentParser) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    content = BytesIO()
    writer.write(content)
    parser.MAX_PDF_PAGES = 1

    with pytest.raises(InputLimitError, match="page limit"):
        parser.parse_bytes("large.pdf", content.getvalue())
